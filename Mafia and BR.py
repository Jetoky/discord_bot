# coding: utf8
from discord import channel
import config
import discord
from typing import List, Dict
from discord.ext import commands
from pathlib import Path
import time
import random
from file_creator import create
from docx import Document
import os
import shutil

"""
Импортируем необходимые модули для работы с помощью библиотеки pip 
"""

bot = commands.Bot(command_prefix='!')
ADMINROLES = "Админ"
SHUFFLED = []
users_bunker = []
users_mafia = []
roles = []

path = Path.cwd()
direct = os.listdir(f"{path}")
try:
    os.makedirs(f"{path}/bunker")
    b = f"{path}/bunker"
except:
    b = f"{path}/bunker"
try:
    os.makedirs(f"{path}/mafia")
    m = f"{path}/mafia"
except:
    m = f"{path}/mafia"
"""
path - получает директрорию работы бота
direct - раскладывает ее на список
"""
async def is_admin(author: discord.Member) -> bool:
    rolenames = [i.name for i in author.roles]
    for i in rolenames:
        if i == ADMINROLES:
            return True
    return False
"""
Проверка на админа (проверяется информация о пользователе, отправившего сообщения) 
"""
@bot.event
async def on_ready():
    print('Logged in as')
    print(bot.user.name)
    print('------')
"""
Функция запускает бота, проверяем его работоспособность с помощью вывода текста 
"""
@bot.command(Pass_context=True, name='FAQ')
async def FAQ(ctx):
    await ctx.send(
        'FAQ:\n1) Для выбора роли напиши !Lobby {Number}(это нужно сделать перед началом игры)\n2)Для смены лобби напиши !Change {Number}\n3)Для начала игры в бункер напиши !play_bunker в чате лобби в котором находишься. Перед игрой необходимо зарегестрироваться при помощи команды !join_bunker. Для завершения игры используйте команду !bunker_stop.\n4)Для начала игры в мафию напиши !play_mafia. Перед началом игры необходимо зарегестрироваться при помощи команды !join_mafia и добавить роли при помощи команды !roles_add. Ведущим будет тот человек, который запустил игру. Для завершения игры введите !mafia_stop.\n5)Для просмотра зарегестрировавшихся игроков введите команду !party_bunker/!party_mafia.\n6)Для просмотра ролей, добавленных в мафию, введите !roles\n7)Для добавления новых ролей в мафию введите !roles_add {Название роли} {Количество ролей}\n8)Для удаления роли введите !roles_del {Название роли}\n9)Для того, что бы очистить список ролей в мафии введите !roles_clear\n10)Для того что бы очистить списко игроков используйте !users_bunker_clear/!users_mafia_clear.')

"""
Функция выводит информацию о сервере в ответ на сообщение пользователя
ctx - в переменную поступает вся информация о юзере
"""
@bot.command(Pass_context=True, name='Lobby')
async def Lobby(ctx, arg):
    member = ctx.message.author
    role_1 = member.guild.get_role(920782582155718687)
    role_2 = member.guild.get_role(920782624681758760)
    role_3 = member.guild.get_role(920782656449421322)
    role_4 = member.guild.get_role(920838063280758784)
    await member.remove_roles(role_1)
    await member.remove_roles(role_2)
    await member.remove_roles(role_3)
    await member.remove_roles(role_4)

    if arg == '1':
        await ctx.send('Подключение к лобби 1.')
        time.sleep(1)
        await ctx.send('Успешно подключен к Лобби 1')
        role_1 = member.guild.get_role(920782582155718687)
        await member.add_roles(role_1)

    if arg == '2':
        await ctx.send('Подключение к лобби 2.')
        time.sleep(1)
        await ctx.send('Успешно подключен к Лобби_2')
        role_2 = member.guild.get_role(920782624681758760)
        await member.add_roles(role_2)

    if arg == '3':
        await ctx.send('Подключение к лобби 3.')
        time.sleep(1)
        await ctx.send('Успешно подключен к Лобби_3')
        role_3 = member.guild.get_role(920782656449421322)
        await member.add_roles(role_3)

    if arg == '4':
        await ctx.send('Подключение к лобби 4.')
        time.sleep(1)
        await ctx.send('Успешно подключен к Лобби_4')
        role_4 = member.guild.get_role(920782656449421322)
        await member.add_roles(role_4)
"""
Функция присваивает пользователю роль, подключая к выбранному лобби, в зависимости от полученного значения переменной
arg 
на вход поступает ctx - информация о пользователе
arg - переменная для выбора лобби
В ответ на команду, пользователь получает сообщения об успешном подключении к лобби
команды time.sleep() - необходимы чтобы бот не спамил сообщениями, а отправлял их с небольшой задержкой
"""
@bot.command(Pass_context=True, name="roles_add")
async def add_role(ctx, name, k):
    for i in range(int(k)):
        roles.append(name)
    await ctx.send(f"Роль \"{name}\" успешно добавлена!")
"""
Функция добавляет новые роли по запросу пользователя для игры в мафию
На вход поступает информация: ctx - инфо о пользователе
name - название роли
k - количество ролей
"""
@bot.command(Pass_context=True, name="roles_del")
async def del_role(ctx, name):
    if name not in roles:
        await ctx.send(f"Роль с именем \"{name}\" не существует.")
        return
    roles.remove(name)
    await ctx.send(f"Роль с именем \"{name}\" успешно удалена.")
"""
Функция проверяет и удаляет роли
на вход поступает: ctx- инфо пользователя
name - необходимая роль
В проверке условия просматривается существует ли указанная роль
Далее если роль найдена - она удаляется, выводя сообщение о результате
"""
@bot.command(Pass_context=True, name='roles_clear')
async def roles_clear(ctx):
    if roles == []:
        await ctx.send('Список ролей пуст!')
        return
    roles.clear()
    await ctx.send('Список ролей успешно очищен!')
"""
Функция проверяет, есть ли в массиве ролей сами роли, в случае их наличия - команда удаляет роли 
На вход: ctx - инфо пользователя
"""
@bot.command(Pass_context=True, name="roles_return")
async def roles_return(ctx):
    await ctx.send(f'Список ролей:\n{roles}')
"""
Функция выводит весь список присутствующих ролей
На вход: ctx - инфо пользователя
"""
@bot.command(Pass_context=True, name="users_mafia_clear")
async def users_mafia_clear(ctx):
    if users_mafia == []:
        await ctx.send('Список игроков пуст!')
        return
    users_mafia.clear()
    await ctx.send('Список игроков в мафию успешно очищен!')
"""
Функция проверяет наличие ролей в игре "Мафия"
На вход: ctx - инфо пользователя
Если массив пустой - выводит пользователю сообщение об этом
Если нет - то удаляет роли и выводит сообщение
"""
@bot.command(Pass_context=True, name="users_bunker_clear")
async def users_bunker_clear(ctx):
    if users_bunker == []:
        await ctx.send('Список игроков пуст!')
        return
    users_bunker.clear()
    await ctx.send('Список игроков в бункер успешно очищен!')
"""
Функция проверяет наличие ролей в игре "Бункер"
На вход: ctx - инфо пользователя
Если массив пустой - выводит пользователю сообщение об этом
Если нет - то удаляет роли и выводит сообщение
"""
@bot.command(Pass_context=True, name='mafia_stop')
async def mafia_stop(ctx):
    shutil.rmtree(f"{path}/mafia/")
    roles.clear()
    await ctx.send('Список ролей успешно очищен!')
    users_mafia.clear()
    await ctx.send('Список игроков в мафию успешно очищен!')
    await ctx.send('Игра в мафию успешно завершена!')
"""
Функция Завершает игру "Мафия"
На вход: ctx - инфо пользователя
После завершения игры удаляется весь массив ролей
Бот выводит сообщение пользователю об изменениях
"""
@bot.command(Pass_context=True, name='bunker_stop')
async def bunker_stop(ctx):
    shutil.rmtree(f"{path}/bunker/")
    users_bunker.clear()
    await ctx.send('Список игроков в бункер успешно очищен!')
"""
Функция завершает игру "Бункер" и очищает список игроков
На вход: ctx - инфо пользователя
"""
@bot.command(Pass_context=True, name='Change')
async def remove_roles(ctx, arg):
    member = ctx.message.author
    role_1 = member.guild.get_role(920782582155718687)
    role_2 = member.guild.get_role(920782624681758760)
    role_3 = member.guild.get_role(920782656449421322)
    role_4 = member.guild.get_role(920838063280758784)
    await member.remove_roles(role_1)
    await member.remove_roles(role_2)
    await member.remove_roles(role_3)
    await member.remove_roles(role_4)

    if arg == '1':
        await ctx.send('Отключение от Лобби...')
        time.sleep(1)
        role_1 = member.guild.get_role(920782582155718687)  # id роли которую будет получать юзер
        await member.add_roles(role_1)
        await ctx.send('Успешно подключен к Лобби_' + str(arg) + '.')

    if arg == '2':
        await ctx.send('Отключение от Лобби...')
        time.sleep(1)
        role_1 = member.guild.get_role(920782624681758760)  # id роли которую будет получать юзер
        await member.add_roles(role_1)
        await ctx.send('Успешно подключен к Лобби_' + str(arg) + '.')

    if arg == '3':
        await ctx.send('Отключение от Лобби...')
        role_1 = member.guild.get_role(920782656449421322)  # id роли которую будет получать юзер
        await member.add_roles(role_1)
        await ctx.send('Успешно подключен к Лобби_' + str(arg) + '.')

    if arg == '4':
        await ctx.send('Отключение от Лобби...')
        role_1 = member.guild.get_role(920782656449421322)  # id роли которую будет получать юзер
        await member.add_roles(role_1)
        await ctx.send('Успешно подключен к Лобби_' + str(arg) + '.')
"""
Функция меняет роль пользователю
На вход: ctx - инфо пользователя, arg - переменная с данными о новом лобби
Проверки условия от 1-4 (в зависимости от введенного аргумента arg) отключают пользователя от первоначального лобби
и подключают к другому, выводя сообщения об изменениях
"""
@bot.command(name="party_bunker")
async def party_bunker(ctx):
    l = users_bunker
    if not l:
        await ctx.send("Игроков нет")
        return
    content = "Список игроков:\n" + "\n".join(
        [f"<@{user}>" for user in l]
    )
    await ctx.send(content)
"""
Функция выводит список игроков
На вход: ctx - инфо пользователя
При проверке условия просматривается наличие игроков в "Бункере"
"""
@bot.command(name="party_mafia")
async def party_mafia(ctx):
    l = users_mafia
    if not l:
        await ctx.send("Игроков нет")
        return
    content = "Список игроков:\n" + "\n".join(
        [f"<@{user}>" for user in l]
    )
    await ctx.send(content)
"""
Функция проверяет наличие игроков в игре "Мафия"
На вход: ctx - инфо пользователя
В положительном прохождения проверки условия варианте выводится список игроков
"""
@bot.command(name="join_bunker")
async def join_bunker(ctx):
    """if ctx.author.id in users_bunker:
        await ctx.send("Вы уже зарегистрированы.")
        return
    else:"""
    users_bunker.append(ctx.author.id)
    number_bunker = len(users_bunker)
    await ctx.send(f"Вы зарегистрировались на игру. Ваш номер {number_bunker}")
"""
Функция регистрирует пользователя на игру "Бункер"
На вход: ctx - инфо пользователя
"""
@bot.command(name="join_mafia")
async def join_mafia(ctx):
    """if ctx.author.id in users_mafia:
        await ctx.send("Вы уже зарегистрированы.")
        return
    else:"""
    users_mafia.append(ctx.author.id)
    number_mafia = len(users_mafia)
    await ctx.send(f"Вы зарегистрировались на игру. Ваш номер {number_mafia}")
"""
Функция регистрирует пользователя на игру "Мафия"
На вход: ctx - инфо пользователя
"""
@bot.command(Pass_context=True)
async def play_bunker(ctx):
    await ctx.send(
        'Да начнётся игра! Сейчас будут отправлены файлы с вашими игровыми персонажами.(Для более интересной игры не подсматривайте)')

    h = 1
    while h < len(users_bunker) + 1:
        time.sleep(0.5)
        create()
        with open(f"{b}/{h}.docx", 'w') as document_bunker:
            document_bunker = Document()
            myfile = open(f'{path}/BabyFile.txt').read()
            document_bunker.add_paragraph(myfile)
            document_bunker.save(f'{b}/' + str(h) + '.docx')
        j = str(h) + '.docx'
        await ctx.send(file=discord.File(f"{b}/{j}"))
        h += 1

"""
Функция запускает игру "Бункер" и выдает пользователю сообщение о запуске игры.
На вход: ctx - инфо пользователя
В цикле для каждого пользователя создается и высылается отдельный файл с необходимой информацией для игры
"""
@bot.command(Pass_context=True)
async def play_mafia(ctx):
    if is_admin == False:
        await ctx.send('Игру может запустить только админ.')
        time.sleep(0.5)
    if len(users_mafia) <= 0:
        await ctx.send('Игроков слишком мало для начала игры!')
        time.sleep(0.5)
        return
    if len(roles) == 0:
        await ctx.send('Роли не добавлены, добавьте их через команду *!add_role* ')
        time.sleep(0.5)
        return
    if len(users_mafia) != len(roles):
        await ctx.send('Количество ролей не совпадает с количеством игроков')
        await ctx.send(party_mafia)
        await ctx.send(roles_return)
        time.sleep(0.5)
        return
    random.shuffle(roles)
    h = 1
    while h < len(users_mafia) + 1:
        with open(f"{m}/{h}.docx", 'w') as document_mafia:
            document_mafia = Document()
            document_mafia.add_paragraph(str(roles[h - 1]))
            document_mafia.save(f'{m}/' + str(h) + '.docx')
            j = str(h) + '.docx'
            await ctx.send(file=discord.File(f"{m}/{j}"))
            h += 1
    await ctx.send('Начнем поиски мафии!')

"""
Функция запускает игру "Мафия" и выдает пользователю сообщение о запуске игры.
На вход: ctx - инфо пользователя
Первая проверка условия - админ ли юзер
Вторая проверка условия - подходит ли количество человек для игры
Третья проверка условия - есть ли роли для игры
Четвертая проверка условия - верно ли количество ролей и участников
В цикле для каждого пользователя создается и высылается отдельный файл с необходимой информацией для игры
"""
bot.run(config.TOKEN)
